#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析单词表文件的实际格式
"""

from docx import Document

def analyze_file(file_path):
    """分析文件内容格式"""
    doc = Document(file_path)

    print(f"\n{'='*60}")
    print(f"文件: {file_path}")
    print(f"{'='*60}")

    # 显示所有段落
    for i, para in enumerate(doc.paragraphs[:50], 1):
        text = para.text.strip()
        if text:
            print(f"{i:3d}. {text}")

    print(f"\n总段落数: {len(doc.paragraphs)}")

# 分析所有文件
files = [
    '/workspace/projects/assets/【上海】六上英语单词背默.docx',
    '/workspace/projects/assets/【上海】八上英语单词背默.docx',
]

for file_path in files:
    analyze_file(file_path)
